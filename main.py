from docx import Document
import json

template = 'input/template.docx'
doc = Document(template)
keyword = None


with open(f'output/{keyword}.json', 'r', encoding='utf-8') as f:
    ep_data = json.load(f)

def fill_docx(fill_data):
    # 文本处理
    for para in doc.paragraphs:
        for key, value in fill_data.items():
            if key in para.text:
                for run in para.runs:
                    run.text = run.text.replace(key, value)

    # 表格处理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in fill_data.items():
                        if key in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(key, value)

fill_docx(ep_data)
doc.save(f'output/XX支行关于{keyword}的贷款申请调查报告.docx')
