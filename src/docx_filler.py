import json
import re
from docx import Document


def get_nested_value(item, keys):
    """
    递归获取 item 中嵌套键列表 keys 的值。
    如果在遍历过程中遇到列表，则对列表内每个字典提取对应键的值并拼接。
    """
    if not keys:
        return item
    key = keys[0]
    if isinstance(item, list):
        results = []
        for subitem in item:
            if isinstance(subitem, dict):
                result = get_nested_value(subitem, keys)
                if result != "":
                    results.append(result)
        return "\n".join(str(r) for r in results)
    elif isinstance(item, dict):
        return get_nested_value(item.get(key, ""), keys[1:])
    else:
        return ""


def replace_placeholder(text, json_data):
    """
    替换给定文本中的占位符：
      - 对于简单键，直接替换；
      - 对于数组形式的键，支持多层嵌套形式，占位符形如 key[sub_key1][sub_key2]；
      - 对于字典形式的值，支持形如 key{sub_key} 的占位符
    """
    # 处理数组形式占位符，支持多层嵌套，如 key[sub_key1][sub_key2]
    for key, value in json_data.items():
        if isinstance(value, list):
            pattern = re.compile(
                r'{}((?:\[[^\]]+\])+)' .format(re.escape(key)))

            def array_replacer(match):
                nested_str = match.group(1)  # 如 "[main_experience][position]"
                keys = re.findall(r'\[([^\]]+)\]', nested_str)
                replaced_value = "\n".join(
                    str(get_nested_value(item, keys)) for item in value
                )
                return replaced_value
            text = pattern.sub(array_replacer, text)

    # 处理字典形式占位符: key{sub_key}
    for key, value in json_data.items():
        if isinstance(value, dict):
            pattern = re.compile(
                r'{}{{([a-zA-Z0-9_]+)}}'.format(re.escape(key)))

            def dict_replacer(match):
                sub_key = match.group(1)
                return str(value.get(sub_key, ""))
            text = pattern.sub(dict_replacer, text)

    # 处理简单形式占位符: key
    for key, value in json_data.items():
        if isinstance(value, (str, int, float)):
            text = text.replace(key, str(value))
    return text


def process_paragraph(paragraph, json_data):
    """
    将一个段落的所有 runs 合并处理，再重新设置回段落，
    避免因占位符被拆分而导致替换失败的问题。
    """
    full_text = "".join(run.text for run in paragraph.runs)
    replaced_text = replace_placeholder(full_text, json_data)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced_text
    else:
        paragraph.add_run(replaced_text)


def process_cell(cell, json_data):
    """
    处理单元格中的段落，并递归处理单元格内可能存在的嵌套表格。
    """
    for paragraph in cell.paragraphs:
        process_paragraph(paragraph, json_data)
    # 如果单元格中存在嵌套表格，则递归处理
    if hasattr(cell, "tables"):
        for table in cell.tables:
            process_table(table, json_data)


def process_table(table, json_data):
    """
    遍历表格中的每个单元格，处理其中的段落和嵌套表格。
    """
    for row in table.rows:
        for cell in row.cells:
            process_cell(cell, json_data)


def process_docx(docx_path, json_data, output_path):
    doc = Document(docx_path)

    # 处理所有段落
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, json_data)

    # 处理所有一级表格（以及递归嵌套的表格）
    for table in doc.tables:
        process_table(table, json_data)

    doc.save(output_path)


def generate_report(file_name):
    json_path = f"output/{file_name}.json"
    docx_path = "input/template.docx"
    output_path = f"output/XX支行关于{file_name}的贷款申请报告.docx"

    with open(json_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    process_docx(docx_path, json_data, output_path)
    print("[Output] 填充完成，输出文件保存至：", output_path)


# 底层运行逻辑，测试用
if __name__ == "__main__":
    try:
        keyword = input("文件名 (无后缀) > ").strip()
        if not keyword:
            raise ValueError("[Error] 文件名不能为空！")
        generate_report(keyword)
    except Exception as e:
        print("[Error] 程序出现错误:", str(e))
