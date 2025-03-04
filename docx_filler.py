import os
import json
from docx import Document


def load_json(json_path):
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"[Error] JSON 文件不存在: {json_path}")
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"[Error] 解析 JSON 文件出错 {json_path}: {e}")
    return data


def fill_docx(template_path, fill_data, output_path):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"[Error] 模板文件不存在: {template_path}")
    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"[Error]打开模板文件出错: {e}")

    if not isinstance(fill_data, dict):
        raise TypeError("[Error] fill_data 必须为字典类型。")

    # 文本段落处理
    for para in doc.paragraphs:
        for key, value in fill_data.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    # 表格处理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in fill_data.items():
                        if key in para.text:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(
                                        key, str(value))
    try:
        doc.save(output_path)
    except Exception as e:
        raise IOError(f"[Error] 保存文档到 {output_path} 出错: {e}")


def generate_report(keyword, template='input/template.docx', output_dir='output'):
    if not keyword:
        raise ValueError("[Error] 企业名称不能为空，你是怎么做到的？")
    json_path = os.path.join(output_dir, f"{keyword}.json")
    fill_data = load_json(json_path)

    output_filename = f"关于{keyword}的贷款申请调查报告.docx"
    output_path = os.path.join(output_dir, output_filename)
    fill_docx(template, fill_data, output_path)

    return output_path
